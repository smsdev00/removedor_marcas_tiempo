import tkinter as tk
from tkinter import messagebox, filedialog
import pyperclip
import re  # Import the re module for regular expressions
import os

# Check if python-docx is installed
try:
    import docx
except ImportError:
    raise ImportError("The 'python-docx' module is required. Please install it using 'pip install python-docx'.")

def delete_text_area_content():
    text_area.delete("1.0", "end")  # Clear text input area button

def paste_text_area_content():
    clipboard_text = pyperclip.paste()
    text_area.insert("1.0", clipboard_text)

def remove_timestamps_handler():
    text = text_area.get("1.0", "end-1c").strip()  # Get text from text area and remove leading/trailing whitespace
    try:
        processed_text = remove_timestamps(text)
        pyperclip.copy(processed_text)  # Copy processed text to clipboard
        text_area.delete("1.0", "end")  # Clear text input area after success
        text_area.insert("1.0", processed_text)  # Fill text area with modified text
        messagebox.showinfo("Success", "Texto copiado al portapapeles\n\n")
    except Exception as e:
        messagebox.showerror("Error", "An error occurred:\n\n" + str(e))

def remove_timestamps(text):
    
    # Use regular expression to remove text inside square brackets
    processed_text = re.sub(r'\[\d{2}:\d{2}\.\d{3} --&gt; \d{2}:\d{2}\.\d{3}\]', '', text)

    # Remove leading and trailing white spaces from each line
    processed_text = '\n'.join([line.strip() for line in processed_text.split('\n')])
    return processed_text.strip()

def select_file():
    file_path = filedialog.askopenfilename(initialdir=os.path.expanduser('~/Downloads'), title="Select .docx file", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
    if file_path:
        try:
            # Read the .docx file
            doc = docx.Document(file_path)
            
            # Extract text from all paragraphs
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Remove text inside brackets
            processed_text = remove_timestamps(text)
            
            pyperclip.copy(processed_text)  # Copy processed text to clipboard
            text_area.delete("1.0", "end")  # Clear text input area after success
            text_area.insert("1.0", processed_text)  # Fill text area with modified text

            # Create a new document
            new_doc = docx.Document()
            
            # Add processed text as a new paragraph
            new_doc.add_paragraph(processed_text)
            
            # Get the directory and file name without extension
            directory, filename = os.path.split(file_path)
            filename_no_ext, ext = os.path.splitext(filename)
            
            # Create new file name with "_removed_timestamp" appended
            new_filename = filename_no_ext + "_removed_timestamp" + ext
            
            # Construct the path for the new file
            new_file_path = os.path.join(directory, new_filename)
            
            # Save the new document with the modified text
            new_doc.save(new_file_path)
            
            # Inform the user about successful operation
            new_file_path = os.path.join(directory, new_filename)
            messagebox.showinfo("Success", f"Texto en el archivo '{new_filename}' modificado y guardado con éxito en '{new_file_path}'.")
        except Exception as e:
            messagebox.showerror("Error", "Ocurrió un error al procesar el archivo .docx:\n\n" + str(e))

def show_tutorial():
    tutorial_text = """\
    Bienvenido al Tutorial de Uso:
    BOTONES:
        * Remover Marcas de Tiempo: Este botón eliminará cualquier marca de tiempo dentro del texto en la ventana de entrada.

        * Limpiar ventana: Este botón limpiará el área de texto en la ventana.

        * Pegar texto desde portapapeles: Este botón pegará el texto actualmente copiado en el portapapeles en la ventana de entrada.

        * Seleccionar archivo .docx: Este botón le permitirá seleccionar un archivo .docx para procesar. 
        Al seleccionar un archivo automaticamente se quitaran las marcas de tiempo del contenido. 
        Se guardara el texto modificado en nuevo archivo, en el mismo directorio del original.
        Con el mismo nombre precedido de '_removed_timestamp'.
    """
    text_area.delete("1.0", "end")  # Clear text input area after success
    text_area.insert("1.0", tutorial_text)  # Fill text area with modified text


root = tk.Tk()
root.title("REMOVEDOR MARCAS DE TIEMPO")

text_area = tk.Text(root, height=20, width=100)
text_area.pack(padx=20, pady=20)

button_remove_timestamps = tk.Button(root, text="Remover Marcas de tiempo", command=remove_timestamps_handler)
button_remove_timestamps.pack(side="left", padx=5, pady=5)

button_delete_text_area_content = tk.Button(root, text="Limpiar ventana", command=delete_text_area_content)
button_delete_text_area_content.pack(side="left", padx=5, pady=5)

button_paste_text_area_content = tk.Button(root, text="Pegar texto desde portapapeles", command=paste_text_area_content)
button_paste_text_area_content.pack(side="left", padx=5, pady=5)

button_select_file = tk.Button(root, text="Seleccionar archivo .docx", command=select_file)
button_select_file.pack(side="left", padx=5, pady=5)

button_tutorial = tk.Button(root, text="Tutorial de Uso", command=show_tutorial)
button_tutorial.pack(side="left", padx=5, pady=5)

root.mainloop()
